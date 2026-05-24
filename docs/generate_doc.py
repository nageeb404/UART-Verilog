#!/usr/bin/env python3
"""
generate_doc.py  —  Creates UART_Documentation.docx
Run: python generate_doc.py
"""

import os, io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE        = os.path.dirname(os.path.abspath(__file__))   # docs/
FIGDIR      = os.path.join(BASE, 'figures')
SCREENDIR   = os.path.join(BASE, 'screenshots')
FIGPATHS    = {}

# ── Word colours ─────────────────────────────────────────────────
BLUE   = RGBColor(26,  82, 118)
WHITE  = RGBColor(255, 255, 255)
LBLUE  = RGBColor(214, 234, 248)

# ── Matplotlib colours ────────────────────────────────────────────
CS  = '#d6eaf8'   # state fill
CSB = '#1a5276'   # state border
CD  = '#fef9e7'   # diamond fill
CDB = '#9a7d0a'   # diamond border
COP = '#e8f8f5'   # operation fill
COB = '#1a8a4a'   # operation border

# ══════════════════════════════════════════════════════════════════
#  MATPLOTLIB HELPERS
# ══════════════════════════════════════════════════════════════════

def savefig(fig, name):
    os.makedirs(FIGDIR, exist_ok=True)
    p = os.path.join(FIGDIR, f"_fig_{name}.png")
    fig.savefig(p, dpi=160, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    FIGPATHS[name] = p
    return p

def state_box(ax, cx, cy, name, out='', w=2.8, h=0.95):
    ax.add_patch(FancyBboxPatch(
        (cx - w/2, cy - h/2), w, h,
        boxstyle='round,pad=0.09', linewidth=1.6,
        linestyle='--', edgecolor=CSB, facecolor=CS, zorder=3))
    yo = 0.18 if out else 0
    ax.text(cx, cy + yo, name, ha='center', va='center',
            fontsize=10, fontweight='bold', color=CSB, zorder=4)
    if out:
        ax.text(cx, cy - 0.17, out, ha='center', va='center',
                fontsize=8.5, color='#2c3e50', style='italic', zorder=4)

def diamond(ax, cx, cy, text, w=2.6, h=0.90):
    pts = [[cx, cy+h/2], [cx+w/2, cy],
           [cx, cy-h/2], [cx-w/2, cy]]
    ax.add_patch(plt.Polygon(pts, closed=True, linewidth=1.2,
                             edgecolor=CDB, facecolor=CD, zorder=3))
    ax.text(cx, cy, text, ha='center', va='center',
            fontsize=8, zorder=4)

def op_box(ax, cx, cy, text, w=2.2, h=0.42):
    ax.add_patch(FancyBboxPatch(
        (cx - w/2, cy - h/2), w, h,
        boxstyle='round,pad=0.05', linewidth=1,
        edgecolor=COB, facecolor=COP, zorder=3))
    ax.text(cx, cy, text, ha='center', va='center',
            fontsize=7.8, zorder=4)

def arr(ax, x1, y1, x2, y2, lbl='', lx=None, ly=None, lha='left', col='#2c3e50'):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=col, lw=1.2), zorder=2)
    if lbl:
        ax.text(lx if lx is not None else (x1+x2)/2 + 0.1,
                ly if ly is not None else (y1+y2)/2,
                lbl, fontsize=8, ha=lha, va='center', color=col)

def side_label(ax, x, y, text, ha='right'):
    ax.text(x, y, text, fontsize=8, ha=ha, va='center', color='#2c3e50')

# ══════════════════════════════════════════════════════════════════
#  FIGURE 1 — UART Frame Format
# ══════════════════════════════════════════════════════════════════
def fig_frame():
    fig, ax = plt.subplots(figsize=(9.5, 2.4))
    ax.set_xlim(-0.2, 12.5)
    ax.set_ylim(-0.8, 2.2)
    ax.axis('off')

    fields = [
        ('Idle',           0.7, '#d5d8dc', '#717d7e'),
        ('Start\n(0)',     0.9, '#fadbd8', '#922b21'),
        ('D0 (LSB)',       0.9, '#d6eaf8', '#1a5276'),
        ('D1',             0.9, '#d6eaf8', '#1a5276'),
        ('D2',             0.9, '#d6eaf8', '#1a5276'),
        ('D3',             0.9, '#d6eaf8', '#1a5276'),
        ('D4',             0.9, '#d6eaf8', '#1a5276'),
        ('D5',             0.9, '#d6eaf8', '#1a5276'),
        ('D6',             0.9, '#d6eaf8', '#1a5276'),
        ('D7 (MSB)',        0.9, '#d6eaf8', '#1a5276'),
        ('Parity\n(opt.)', 1.0, '#e8f8f5', '#1a8a4a'),
        ('Stop\n(1)',      0.9, '#fadbd8', '#922b21'),
        ('Idle',           0.7, '#d5d8dc', '#717d7e'),
    ]

    x = 0
    for name, w, fill, edge in fields:
        ax.add_patch(mpatches.FancyBboxPatch(
            (x, 0.4), w, 1.2,
            boxstyle='square,pad=0', linewidth=1.2,
            edgecolor=edge, facecolor=fill))
        ax.text(x + w/2, 1.0, name, ha='center', va='center',
                fontsize=7.8, color=edge, fontweight='bold')
        x += w

    # ── data bits brace ──────────────────────────────────────────
    x0 = 0.7 + 0.9          # start of D0
    x1 = x0 + 8 * 0.9       # end of D7
    ax.annotate('', xy=(x1, 0.15), xytext=(x0, 0.15),
                arrowprops=dict(arrowstyle='<->', color='#1a5276', lw=1.2))
    ax.text((x0+x1)/2, -0.05, '8 Data Bits — LSB first',
            ha='center', va='top', fontsize=8.5, color='#1a5276')

    ax.set_title('UART Frame Structure', fontsize=11, fontweight='bold',
                 color='#1a5276', pad=5)
    fig.tight_layout(pad=0.3)
    return savefig(fig, 'frame')

# ══════════════════════════════════════════════════════════════════
#  FIGURE 2 — 16× Oversampling Timing Diagram
# ══════════════════════════════════════════════════════════════════
def fig_oversampling():
    fig, ax = plt.subplots(figsize=(9.5, 3.5))
    ax.set_xlim(-1, 35)
    ax.set_ylim(-1.8, 6.2)
    ax.axis('off')

    SY, CY, H, CT = 4.2, 2.0, 1.0, 0.45

    # 16× clock signal
    for t in range(33):
        ax.plot([t, t, t+CT, t+CT, t+1], [CY, CY+H, CY+H, CY, CY],
                color='#1a5276', lw=1.1)
    ax.text(-0.8, CY+H/2, 'UART\n16× CLK', ha='right', va='center',
            fontsize=8.5, color='#1a5276')

    # RX serial signal (idle=1 → start=0 → D0=1)
    rx = [(-0.5, H), (0, H), (0, 0), (16, 0), (16, H), (34, H)]
    ax.plot([p[0] for p in rx], [SY+p[1] for p in rx],
            color='#2c3e50', lw=2.2)
    ax.text(-0.8, SY+H/2, 'RX\n(sync)', ha='right', va='center',
            fontsize=8.5, color='#2c3e50')

    # Bit period arrows
    ax.annotate('', xy=(16, SY+H+0.45), xytext=(0, SY+H+0.45),
                arrowprops=dict(arrowstyle='<->', color='#922b21', lw=1.3))
    ax.text(8, SY+H+0.65, 'Start Bit', ha='center', fontsize=9, color='#922b21',
            fontweight='bold')

    ax.annotate('', xy=(32, SY+H+0.45), xytext=(16, SY+H+0.45),
                arrowprops=dict(arrowstyle='<->', color='#1a5276', lw=1.3))
    ax.text(24, SY+H+0.65, 'D0 Bit', ha='center', fontsize=9, color='#1a5276',
            fontweight='bold')

    # Tick count brackets
    ax.annotate('', xy=(8, CY-0.55), xytext=(0, CY-0.55),
                arrowprops=dict(arrowstyle='<->', color='#922b21', lw=1.0))
    ax.text(4, CY-0.85, '8 ticks', ha='center', fontsize=8, color='#922b21')

    ax.annotate('', xy=(32, CY-0.55), xytext=(16, CY-0.55),
                arrowprops=dict(arrowstyle='<->', color='#1a5276', lw=1.0))
    ax.text(24, CY-0.85, '16 ticks', ha='center', fontsize=8, color='#1a5276')

    # Sampling markers
    for xm, col, lbl, ly in [(8, '#922b21', 'START\n(sample @ tick 7)', SY-0.55),
                              (24, '#1a5276', 'MIDBIT\n(sample @ tick 15)', SY-0.55)]:
        ax.axvline(xm, ymin=(CY-0.2)/7.5, ymax=(SY+H+0.3)/7.5,
                   color=col, lw=1.8, linestyle=':', alpha=0.85, zorder=4)
        ax.text(xm, ly, lbl, ha='center', va='top', fontsize=8,
                color=col, fontweight='bold')

    ax.set_title('16× Oversampling — Start Bit and First Data Bit',
                 fontsize=11, fontweight='bold', color='#1a5276', pad=6)
    fig.tight_layout(pad=0.5)
    return savefig(fig, 'oversampling')

# ══════════════════════════════════════════════════════════════════
#  FIGURE 3 — 2-FF Synchronizer Block Diagram
# ══════════════════════════════════════════════════════════════════
def fig_2ff():
    fig, ax = plt.subplots(figsize=(8.5, 3.4))
    ax.set_xlim(0, 11)
    ax.set_ylim(0.2, 4.2)
    ax.axis('off')

    # Clock domain shading
    ax.add_patch(FancyBboxPatch((2.5, 0.4), 7.5, 2.8,
                                boxstyle='round,pad=0.1', linewidth=1.5,
                                linestyle='--', edgecolor='#1a5276',
                                facecolor='#eaf0f9', zorder=1))
    ax.text(6.25, 3.55, 'System Clock Domain  (clk)',
            ha='center', va='center', fontsize=9, color='#1a5276',
            style='italic')

    # FF1
    ax.add_patch(FancyBboxPatch((3.2, 1.0), 1.6, 1.6,
                                boxstyle='round,pad=0.05', linewidth=1.4,
                                edgecolor='#1a5276', facecolor='#d6eaf8', zorder=3))
    ax.text(4.0, 1.8, 'FF1\n(rx_meta)', ha='center', va='center',
            fontsize=9, color='#1a5276', fontweight='bold')

    # FF2
    ax.add_patch(FancyBboxPatch((6.2, 1.0), 1.6, 1.6,
                                boxstyle='round,pad=0.05', linewidth=1.4,
                                edgecolor='#1a5276', facecolor='#d6eaf8', zorder=3))
    ax.text(7.0, 1.8, 'FF2\n(rx_sync)', ha='center', va='center',
            fontsize=9, color='#1a5276', fontweight='bold')

    # Arrows & labels
    arr(ax, 0.3, 1.8, 3.2, 1.8)
    ax.text(0.3, 1.8, 'rx\n(async)', ha='right', va='center',
            fontsize=9, color='#2c3e50')

    arr(ax, 4.8, 1.8, 6.2, 1.8)
    ax.text(5.5, 2.05, 'may be\nmetastable', ha='center', va='bottom',
            fontsize=7.5, color='#922b21', style='italic')

    arr(ax, 7.8, 1.8, 9.8, 1.8)
    ax.text(9.85, 1.8, 'rx_sync\n(stable → FSM)', ha='left', va='center',
            fontsize=9, color='#1a8a4a', fontweight='bold')

    # Clock edges at bottom of FFs
    for x in [4.0, 7.0]:
        ax.annotate('', xy=(x, 1.0), xytext=(x - 0.25, 1.25),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.0))

    ax.text(5.25, 0.5, 'clk ↑', ha='center', va='center',
            fontsize=8, color='#555', style='italic')

    ax.set_title('2-FF Metastability Synchronizer on RX Input',
                 fontsize=11, fontweight='bold', color='#1a5276', pad=6)
    fig.tight_layout(pad=0.3)
    return savefig(fig, '2ff')

# ══════════════════════════════════════════════════════════════════
#  FIGURE 4 — Top-Level Block Diagram
# ══════════════════════════════════════════════════════════════════
def fig_toplevel():
    fig, ax = plt.subplots(figsize=(9.5, 6.5))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis('off')

    # ── UART_top boundary ────────────────────────────────────────
    ax.add_patch(FancyBboxPatch((1.8, 0.4), 8.4, 7.2,
                                boxstyle='round,pad=0.12', linewidth=2,
                                linestyle='-', edgecolor='#1a5276',
                                facecolor='#f0f6fc', zorder=1))
    ax.text(6.0, 7.75, 'UART_top', ha='center', va='center',
            fontsize=12, fontweight='bold', color='#1a5276')

    def module(cx, cy, w, h, name, color='#d6eaf8'):
        ax.add_patch(FancyBboxPatch(
            (cx - w/2, cy - h/2), w, h,
            boxstyle='round,pad=0.08', linewidth=1.4,
            edgecolor='#1a5276', facecolor=color, zorder=3))
        ax.text(cx, cy, name, ha='center', va='center',
                fontsize=9.5, fontweight='bold', color='#1a5276', zorder=4)

    # Sub-modules
    module(4.0, 6.2, 2.6, 0.9, 'Baud_Rate_Generator', '#e8f8f5')
    module(4.0, 4.2, 2.6, 1.2, 'UART_tx\n(Transmitter)')
    module(4.0, 1.8, 2.6, 1.2, 'UART_rx\n(Receiver)')

    # s_tick wire (vertical)
    ax.plot([4.0, 4.0], [5.75, 5.35], color='#555', lw=1.4, zorder=2)
    ax.plot([4.0, 4.0], [5.35, 3.8],  color='#555', lw=1.4, linestyle='--', zorder=2)
    ax.plot([4.0, 4.0], [3.5, 2.35],  color='#555', lw=1.4, linestyle='--', zorder=2)
    ax.text(4.15, 4.55, 's_tick', fontsize=8, color='#555', style='italic')

    # ── External ports (left side) ──────────────────────────────
    def ext_port(y, label, arrow_to, right=False):
        xa = 0.2 if not right else 11.8
        xb = arrow_to[0]
        yb = arrow_to[1]
        ax.annotate('', xy=(xb, yb), xytext=(xa, y),
                    arrowprops=dict(arrowstyle='->', color='#2c3e50', lw=1.2), zorder=2)
        ha = 'right' if not right else 'left'
        ax.text(xa + (-0.05 if not right else 0.05), y, label,
                ha=ha, va='center', fontsize=8.5, color='#2c3e50')

    ext_port(7.0, 'clk, reset_n', (1.8, 7.0), right=False)
    ext_port(4.8, 'tx_din[7:0]', (2.7, 4.5), right=False)
    ext_port(4.2, 'tx_start', (2.7, 4.0), right=False)
    ext_port(3.8, 'parity_select[1:0]', (2.7, 3.7), right=False)
    ext_port(1.6, 'rx', (2.7, 1.7), right=False)

    # ── External ports (right side — outputs) ───────────────────
    def out_port(y, label, from_pt):
        ax.annotate('', xy=(11.8, y), xytext=(from_pt[0], from_pt[1]),
                    arrowprops=dict(arrowstyle='->', color='#1a8a4a', lw=1.2), zorder=2)
        ax.text(11.85, y, label, ha='left', va='center',
                fontsize=8.5, color='#1a8a4a')

    out_port(4.6, 'tx', (5.3, 4.5))
    out_port(4.0, 'tx_done_tick', (5.3, 3.9))
    out_port(2.0, 'rx_dout[7:0]', (5.3, 2.0))
    out_port(1.5, 'rx_done_tick', (5.3, 1.6))
    out_port(1.0, 'parity_error', (5.3, 1.2))
    out_port(0.5, 'frame_error',  (5.3, 0.8))

    # clk/reset_n distribution line
    ax.plot([1.8, 1.8], [4.8, 1.2], color='#555', lw=1.0,
            linestyle=':', zorder=2)

    ax.set_title('UART_top — Top-Level Block Diagram',
                 fontsize=11, fontweight='bold', color='#1a5276', pad=8)
    fig.tight_layout(pad=0.5)
    return savefig(fig, 'toplevel')

# ══════════════════════════════════════════════════════════════════
#  FIGURE 5 — TX ASMD Chart  (single-column, no overlaps)
# ══════════════════════════════════════════════════════════════════
def fig_tx_asmd():
    fig, ax = plt.subplots(figsize=(8, 18))
    ax.set_xlim(0, 8)
    ax.set_ylim(-2.2, 18)
    ax.axis('off')

    CX = 4.0   # single centre column

    # ── States ──────────────────────────────────────────────────
    state_box(ax, CX, 17.0, 'idle',   'tx ← 1')
    state_box(ax, CX, 13.7, 'start',  'tx ← 0')
    state_box(ax, CX, 10.5, 'data',   'tx ← b_reg[0]')
    state_box(ax, CX,  4.0, 'parity', 'tx ← p_reg')
    state_box(ax, CX,  0.6, 'stop',   'tx ← 1')

    # ── Condition diamonds ───────────────────────────────────────
    diamond(ax, CX, 15.7, 'tx_start == 1 ?',      w=2.8)
    diamond(ax, CX, 12.4, 's_tick && s == 15 ?',   w=3.0)
    diamond(ax, CX,  8.7, 's_tick && s == 15 ?',   w=3.0)
    diamond(ax, CX,  6.8, 'n == DBITS−1 ?',        w=2.8)
    diamond(ax, CX,  5.5, 'parity_select ≠ 0 ?',  w=2.8)
    diamond(ax, CX,  2.7, 's_tick && s == 15 ?',   w=3.0)
    diamond(ax, CX, -0.7, 's == SB_TICK−1 ?',      w=3.0)

    # ── Datapath operation boxes ─────────────────────────────────
    op_box(ax, CX,  14.7, 's←0  b←tx_din  p←parity(tx_din)', w=3.4)
    op_box(ax, CX,  11.5, 's←0   n←0', w=1.8)
    op_box(ax, CX,   7.8, 's←0   b←b>>1', w=2.2)
    op_box(ax,  6.2,  6.8, 'n←n+1', w=1.2, h=0.38)   # No branch of n==DBITS-1
    op_box(ax, CX,   1.8, 's←0', w=1.0)
    op_box(ax, CX,  -1.5, 'tx_done_tick←1', w=2.2)

    # ── Forward flow arrows ──────────────────────────────────────
    arr(ax, CX, 16.525, CX, 16.15)                       # idle → tx_start?
    arr(ax, CX, 15.25,  CX, 14.91, 'Yes', lx=CX+0.1)    # Yes → op_load
    arr(ax, CX, 14.49,  CX, 14.175)                      # op → start
    arr(ax, CX, 13.225, CX, 12.85)                       # start → s==15?
    arr(ax, CX, 11.95,  CX, 11.71, 'Yes', lx=CX+0.1)    # Yes → op_n0
    arr(ax, CX, 11.29,  CX, 10.975)                      # op → data
    arr(ax, CX, 10.025, CX,  9.15)                       # data → s==15?
    arr(ax, CX,  8.25,  CX,  8.01, 'Yes', lx=CX+0.1)    # Yes → op_shift
    arr(ax, CX,  7.59,  CX,  7.25)                       # op → n==DBITS-1?
    arr(ax, CX,  6.35,  CX,  5.95, 'Yes', lx=CX+0.1)    # Yes → parity_sel?
    arr(ax, CX,  5.05,  CX,  4.475,'Yes', lx=CX+0.1)    # Yes → parity state
    arr(ax, CX,  3.525, CX,  3.15)                       # parity → s==15?
    arr(ax, CX,  2.25,  CX,  2.01, 'Yes', lx=CX+0.1)    # Yes → op_s0
    arr(ax, CX,  1.59,  CX,  1.075)                      # op_s0 → stop
    arr(ax, CX,  0.125, CX, -0.25)                       # stop → s==SB?
    arr(ax, CX, -1.15,  CX, -1.29, 'Yes', lx=CX+0.1)    # Yes → op_done

    # ── Stop done → back to idle (left-side loop) ────────────────
    ax.plot([CX, CX],   [-1.71, -1.9], color='#2c3e50', lw=1.2)
    ax.plot([CX, 0.3],  [-1.9,  -1.9], color='#2c3e50', lw=1.2)
    ax.plot([0.3, 0.3], [-1.9,  17.0], color='#2c3e50', lw=1.2)
    arr(ax, 0.3, 17.0, CX-1.4, 17.0)

    # ── tx_start No → stay in idle (right loop, dashed) ──────────
    ax.plot([CX+1.4, 7.2], [15.7, 15.7], color='#999', lw=1.1, ls=':')
    ax.plot([7.2, 7.2],    [15.7, 17.0], color='#999', lw=1.1, ls=':')
    arr(ax, 7.2, 17.0, CX+1.4, 17.0, col='#999')
    ax.text(7.25, 16.35, 'No', fontsize=8, ha='left', va='center', color='#999')

    # ── Hold labels for s==15 and s==SB No branches ──────────────
    ax.text(CX+1.6, 12.4,  'No\n(hold)', fontsize=7.5, ha='left', va='center', color='#888')
    ax.text(CX+1.6,  8.7,  'No\n(hold)', fontsize=7.5, ha='left', va='center', color='#888')
    ax.text(CX+1.6,  2.7,  'No\n(hold)', fontsize=7.5, ha='left', va='center', color='#888')
    ax.text(CX+1.6, -0.7,  'No\n(hold)', fontsize=7.5, ha='left', va='center', color='#888')

    # ── n==DBITS-1 No → n←n+1 → back to data (right side) ───────
    arr(ax, CX+1.4, 6.8, 5.6, 6.8, 'No', lx=CX+1.45)        # → n←n+1 box
    ax.plot([6.8, 7.2], [6.8, 6.8], color='#2c3e50', lw=1.2)  # right from box
    ax.plot([7.2, 7.2], [6.8, 10.5], color='#2c3e50', lw=1.2) # up
    arr(ax, 7.2, 10.5, CX+1.4, 10.5)                          # into data right

    # ── parity_sel No → stop (right bypass) ──────────────────────
    ax.plot([CX+1.4, 7.5], [5.5, 5.5], color='#2c3e50', lw=1.2)
    ax.plot([7.5, 7.5],    [5.5, 0.6], color='#2c3e50', lw=1.2)
    arr(ax, 7.5, 0.6, CX+1.4, 0.6, 'No', lx=7.1, lha='left')

    ax.set_title('UART Transmitter — ASMD Chart',
                 fontsize=12, fontweight='bold', color='#1a5276', pad=10)
    fig.tight_layout(pad=0.4)
    return savefig(fig, 'tx_asmd')

# ══════════════════════════════════════════════════════════════════
#  FIGURE 6 — RX ASMD Chart  (single-column, no overlaps)
# ══════════════════════════════════════════════════════════════════
def fig_rx_asmd():
    fig, ax = plt.subplots(figsize=(8, 18))
    ax.set_xlim(0, 8)
    ax.set_ylim(-2.2, 18)
    ax.axis('off')

    CX = 4.0

    # ── States ──────────────────────────────────────────────────
    state_box(ax, CX, 17.0, 'idle',
              'rx_done_tick ← 0\nparity_error ← 0   frame_error ← 0',
              w=3.6, h=1.1)
    state_box(ax, CX, 13.7, 'start')
    state_box(ax, CX, 10.5, 'data')
    state_box(ax, CX,  4.0, 'parity')
    state_box(ax, CX,  0.6, 'stop',
              'rx_dout ← b_reg\nrx_done_tick ← 1',
              w=3.0, h=1.1)

    # ── Condition diamonds ───────────────────────────────────────
    diamond(ax, CX, 15.6, 'rx_sync == 0 ?',       w=2.8)
    diamond(ax, CX, 12.4, 's_tick && s == 7 ?',    w=2.8)
    diamond(ax, CX,  8.7, 's_tick && s == 15 ?',   w=3.0)
    diamond(ax, CX,  6.8, 'n == DBITS−1 ?',        w=2.8)
    diamond(ax, CX,  5.5, 'parity_select ≠ 0 ?',  w=2.8)
    diamond(ax, CX,  2.7, 's_tick && s == 15 ?',   w=3.0)
    diamond(ax, CX, -0.7, 's == SB_TICK−1 ?',      w=3.0)

    # ── Datapath operation boxes ─────────────────────────────────
    op_box(ax, CX,  14.7, 's←0', w=1.0)
    op_box(ax, CX,  11.5, 's←0   n←0   b←0', w=2.4)
    op_box(ax, CX,   7.8, 's←0   b←{rx_sync, b[DBITS-1:1]}', w=3.4)
    op_box(ax,  6.2,  6.8, 'n←n+1', w=1.2, h=0.38)
    op_box(ax, CX,   1.8, 'p←rx_sync', w=1.8)
    op_box(ax, CX,  -1.5, 'frame_error←~rx_sync\nparity_error←check(b,p)',
           w=3.4, h=0.52)

    # ── Forward flow arrows ──────────────────────────────────────
    # idle h=1.1 → bottom = 17.0-0.55 = 16.45
    arr(ax, CX, 16.45,  CX, 16.05)                       # idle → rx_sync==0?
    arr(ax, CX, 15.15,  CX, 14.91, 'Yes', lx=CX+0.1)    # Yes → op_s0
    arr(ax, CX, 14.49,  CX, 14.175)                      # op → start
    arr(ax, CX, 13.225, CX, 12.85)                       # start → s==7?
    arr(ax, CX, 11.95,  CX, 11.71, 'Yes', lx=CX+0.1)    # Yes → op_n0
    arr(ax, CX, 11.29,  CX, 10.975)                      # op → data
    arr(ax, CX, 10.025, CX,  9.15)                       # data → s==15?
    arr(ax, CX,  8.25,  CX,  8.06, 'Yes', lx=CX+0.1)    # Yes → op_shift
    arr(ax, CX,  7.54,  CX,  7.25)                       # op → n==DBITS-1?
    arr(ax, CX,  6.35,  CX,  5.95, 'Yes', lx=CX+0.1)    # Yes → parity_sel?
    arr(ax, CX,  5.05,  CX,  4.475,'Yes', lx=CX+0.1)    # Yes → parity state
    arr(ax, CX,  3.525, CX,  3.15)                       # parity → s==15?
    arr(ax, CX,  2.25,  CX,  2.01, 'Yes', lx=CX+0.1)    # Yes → op_p
    # op_p bottom = 1.8-0.21=1.59 ; stop top (h=1.1) = 0.6+0.55=1.15
    arr(ax, CX,  1.59,  CX,  1.15)                       # op_p → stop
    # stop bottom = 0.6-0.55=0.05 ; s==SB top = -0.7+0.45=-0.25
    arr(ax, CX,  0.05,  CX, -0.25)                       # stop → s==SB?
    # op_done h=0.52, top = -1.5+0.26=-1.24 ; s==SB bottom = -1.15
    arr(ax, CX, -1.15,  CX, -1.24, 'Yes', lx=CX+0.1)    # Yes → op_done

    # ── Stop done → back to idle ──────────────────────────────────
    # op_done bottom = -1.5-0.26=-1.76
    ax.plot([CX, CX],   [-1.76, -1.9], color='#2c3e50', lw=1.2)
    ax.plot([CX, 0.3],  [-1.9,  -1.9], color='#2c3e50', lw=1.2)
    ax.plot([0.3, 0.3], [-1.9,  17.0], color='#2c3e50', lw=1.2)
    arr(ax, 0.3, 17.0, CX-1.8, 17.0)

    # ── rx_sync No → stay in idle (right loop, dashed) ───────────
    ax.plot([CX+1.4, 7.2], [15.6, 15.6], color='#999', lw=1.1, ls=':')
    ax.plot([7.2, 7.2],    [15.6, 17.0], color='#999', lw=1.1, ls=':')
    arr(ax, 7.2, 17.0, CX+1.4, 17.0, col='#999')
    ax.text(7.25, 16.3, 'No', fontsize=8, ha='left', va='center', color='#999')

    # ── Hold labels ───────────────────────────────────────────────
    ax.text(CX+1.6, 12.4,  'No\n(hold)', fontsize=7.5, ha='left', va='center', color='#888')
    ax.text(CX+1.6,  8.7,  'No\n(hold)', fontsize=7.5, ha='left', va='center', color='#888')
    ax.text(CX+1.6,  2.7,  'No\n(hold)', fontsize=7.5, ha='left', va='center', color='#888')
    ax.text(CX+1.6, -0.7,  'No\n(hold)', fontsize=7.5, ha='left', va='center', color='#888')

    # ── n==DBITS-1 No → n←n+1 → back to data ────────────────────
    arr(ax, CX+1.4, 6.8, 5.6, 6.8, 'No', lx=CX+1.45)
    ax.plot([6.8, 7.2], [6.8, 6.8], color='#2c3e50', lw=1.2)
    ax.plot([7.2, 7.2], [6.8, 10.5], color='#2c3e50', lw=1.2)
    arr(ax, 7.2, 10.5, CX+1.4, 10.5)

    # ── parity_sel No → stop (right bypass) ──────────────────────
    ax.plot([CX+1.4, 7.5], [5.5, 5.5], color='#2c3e50', lw=1.2)
    ax.plot([7.5, 7.5],    [5.5, 0.6], color='#2c3e50', lw=1.2)
    arr(ax, 7.5, 0.6, CX+1.5, 0.6, 'No', lx=7.1, lha='left')

    ax.set_title('UART Receiver — ASMD Chart',
                 fontsize=12, fontweight='bold', color='#1a5276', pad=10)
    fig.tight_layout(pad=0.4)
    return savefig(fig, 'rx_asmd')

# ══════════════════════════════════════════════════════════════════
#  DOCX HELPERS
# ══════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = BLUE
        run.font.bold = True
    return h

def add_para(doc, text='', size=11, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_figure(doc, img_path, caption, width=Inches(5.8)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(img_path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cr = cap.add_run(caption)
    cr.font.size = Pt(9)
    cr.font.italic = True
    cr.font.color.rgb = RGBColor(80, 80, 80)

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header row
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(9.5)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(c, '1a5276')

    # Data rows
    for j, row in enumerate(rows):
        dr = t.rows[j + 1]
        bg = 'daeaf8' if j % 2 == 0 else 'ffffff'
        for i, val in enumerate(row):
            c = dr.cells[i]
            c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5)
            set_cell_bg(c, bg)

    if col_widths:
        for j, row in enumerate(t.rows):
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def page_break(doc):
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════

def build_doc():
    doc = Document()

    # ── Page layout ──────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width   = Cm(21)
    sec.page_height  = Cm(29.7)
    sec.left_margin  = Cm(2.54)
    sec.right_margin = Cm(2.54)
    sec.top_margin   = Cm(2.54)
    sec.bottom_margin= Cm(2.54)

    # ════════════════════════════════════════════════════════════
    #  TITLE PAGE
    # ════════════════════════════════════════════════════════════
    for _ in range(6): doc.add_paragraph()

    uni = doc.add_paragraph()
    uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ur = uni.add_run('Menofia University — Faculty of Engineering')
    ur.font.size = Pt(13); ur.font.bold = True; ur.font.color.rgb = BLUE

    doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run('Digital IC Design Using Verilog HDL')
    tr.font.size = Pt(16); tr.font.bold = True; tr.font.color.rgb = BLUE

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr2 = t2.add_run(
        'Universal Asynchronous Receiver-Transmitter (UART)\nRTL Implementation')
    tr2.font.size = Pt(22); tr2.font.bold = True

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('— Implemented with Verilog HDL —')
    sr.font.size = Pt(12); sr.font.italic = True

    for _ in range(5): doc.add_paragraph()

    au = doc.add_paragraph()
    au.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aur = au.add_run('Ahmed Ahmed Nageeb Elbermawy')
    aur.font.size = Pt(14); aur.font.bold = True

    for lbl in ['Menofia University, Faculty of Engineering',
                'Tool: Xilinx Vivado  |  Language: Verilog HDL']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(lbl)
        r.font.size = Pt(11)

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════
    add_heading(doc, 'Table of Contents', level=1)
    toc_items = [
        ('1.', 'Introduction', None),
        ('2.', 'Basic Concepts of UART', [
            'Asynchronous Communication', 'UART Frame Format',
            'Baud Rate', 'Parity Bit', 'Stop Bits']),
        ('3.', 'Design Methodology', [
            '16× Oversampling Procedure',
            'Input Synchronization (2-FF Synchronizer)']),
        ('4.', 'Top-Level Design Diagram', None),
        ('5.', 'Component Breakdown', [
            'Baud Rate Generator',
            'UART Transmitter',
            'UART Receiver',
            'Top-Level Wrapper']),
        ('6.', 'ASMD Charts', [
            'ASMD Chart of the UART Transmitter',
            'ASMD Chart of the UART Receiver']),
        ('7.', 'Testbenches & Simulation', [
            'Transmitter Testbench',
            'Full UART Loopback Testbench']),
        ('8.', 'References', None),
    ]
    for num, title, subs in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f'{num}  {title}')
        r.font.size = Pt(11); r.font.bold = True
        if subs:
            for s in subs:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(2)
                sp.paragraph_format.left_indent = Cm(0.8)
                sr2 = sp.add_run(f'—  {s}')
                sr2.font.size = Pt(10.5)
                sr2.font.color.rgb = RGBColor(50, 100, 150)

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    #  1. INTRODUCTION
    # ════════════════════════════════════════════════════════════
    add_heading(doc, '1.  Introduction', 1)

    add_para(doc,
        'UART (Universal Asynchronous Receiver-Transmitter) is one of the most '
        'fundamental serial communication protocols in digital systems. It enables '
        'data exchange between two devices over just two wires — TX (transmit) and '
        'RX (receive) — without requiring a shared clock signal. This makes it '
        'simple, low-cost, and compatible with virtually every microcontroller, '
        'FPGA, and embedded system.')

    add_para(doc,
        'This report documents the design and RTL implementation of a complete '
        'UART transceiver in Verilog HDL using Xilinx Vivado. The design covers '
        'the full UART frame format including start bit, data bits, optional parity, '
        'and stop bit. The receiver employs 16× oversampling to reliably recover '
        'serial data without a shared clock reference. A 2-FF input synchronizer '
        'on the RX line prevents metastability from propagating into the FSM — a '
        'critical requirement in any real-world digital design.')

    add_para(doc,
        'The design follows a modular, hierarchical approach. Each component is '
        'developed and verified independently through simulation testbenches before '
        'integration into the top-level wrapper. Key features of this implementation:')

    for bullet in [
        'Parameterized design — DBITS and SB_TICK are configurable at elaboration time',
        'Runtime-selectable parity — even, odd, or none via a 2-bit parity_select input',
        'Parity error detection — flags a mismatch between received and expected parity',
        'Frame error detection — flags a missing stop bit (stop bit sampled as logic 0)',
        '2-FF input synchronizer on RX — prevents metastability propagation',
        'Clean RTL style — separate present-state and next-state always blocks throughout',
        'Two testbenches — standalone TX verification and full TX→RX loopback '
            'with automated PASS/FAIL output',
    ]:
        add_bullet(doc, bullet)

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    #  2. BASIC CONCEPTS
    # ════════════════════════════════════════════════════════════
    add_heading(doc, '2.  Basic Concepts of UART', 1)

    add_heading(doc, 'Asynchronous Communication', 2)
    add_para(doc,
        'The defining characteristic of UART is its asynchronous operation — the '
        'transmitter and receiver share no common clock. Each device uses its own '
        'internal oscillator. Synchronization is achieved entirely through the frame '
        'structure: a start bit signals the beginning and stop bits mark the end. '
        'Both sides must be pre-configured to the same baud rate and format.')

    add_heading(doc, 'UART Frame Format', 2)
    add_para(doc,
        'A standard UART frame consists of the following fields transmitted serially '
        'in order: an idle high line, a start bit (logic 0), the data bits (LSB first), '
        'an optional parity bit, and one or more stop bits (logic 1). '
        'Figure 1 illustrates this structure.')

    add_figure(doc, FIGPATHS['frame'],
               'Figure 1.  UART Frame Structure',
               width=Inches(6.0))

    add_table(doc,
        ['Field', 'Logic Level', 'Description'],
        [('Idle',       '1',     'Line rests high when no data is being sent'),
         ('Start bit',  '0',     'Always logic low — signals the start of a new frame'),
         ('Data bits',  'D0–D7', 'Sent LSB first; typically 8 bits'),
         ('Parity bit', '0 / 1', 'Optional error-detection bit (even, odd, or omitted)'),
         ('Stop bit',   '1',     'Always logic high — marks the end of the frame')],
        col_widths=[Cm(2.8), Cm(2.8), Cm(9.5)])

    add_heading(doc, 'Baud Rate', 2)
    add_para(doc,
        'The baud rate defines the number of bits transmitted per second (bps). '
        'Both devices must use identical baud rates; a mismatch results in corrupted '
        'data. With a 100 MHz system clock and 16× oversampling, the Baud Rate '
        'Generator counter limit is:')
    eq = doc.add_paragraph()
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq.paragraph_format.space_after = Pt(6)
    eqr = eq.add_run('FINAL_VALUE  =  F_clk  /  (16 × BaudRate)  −  1')
    eqr.font.size = Pt(11); eqr.font.bold = True; eqr.font.italic = True

    add_table(doc,
        ['Baud Rate (bps)', 'FINAL_VALUE  (100 MHz)', 'BRG_BITS needed'],
        [('9 600',   '651', '10'),
         ('19 200',  '325', '9'),
         ('38 400',  '162', '8'),
         ('57 600',  '108', '7'),
         ('115 200', '54',  '6')],
        col_widths=[Cm(4.5), Cm(5.0), Cm(4.0)])

    add_heading(doc, 'Parity Bit', 2)
    add_para(doc,
        'An optional parity bit provides single-bit error detection. In even parity '
        'the parity bit is set so the total number of logic-1 bits (data + parity) is '
        'even; in odd parity the total is odd. If the receiver calculates a different '
        'parity than received, a parity_error flag is asserted.')

    add_heading(doc, 'Stop Bits', 2)
    add_para(doc,
        'Stop bits are always logic high and mark the end of the frame. They provide '
        'the receiver time to process the received byte. This design supports a '
        'configurable stop-bit duration via the SB_TICK parameter '
        '(16 = 1 stop bit, 32 = 2 stop bits).')

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    #  3. DESIGN METHODOLOGY
    # ════════════════════════════════════════════════════════════
    add_heading(doc, '3.  Design Methodology', 1)

    add_heading(doc, 'UART Receiving Subsystem', 2)
    add_para(doc,
        'In UART communication the transmitted signal does not include a clock, '
        'meaning the receiver has no explicit timing reference to know exactly when '
        'each bit begins and ends. Instead, both transmitter and receiver rely on '
        'a pre-agreed baud rate to stay synchronised in time.')

    add_heading(doc, '16× Oversampling Procedure', 2)
    add_para(doc,
        'To accurately recover the transmitted data, the receiver uses an '
        'oversampling technique. The Baud Rate Generator produces a sampling tick '
        '(s_tick) at 16 times the baud rate. The receiver counts these ticks to '
        'locate the centre of each bit — the point where the signal is most stable '
        'and least susceptible to noise or jitter.')

    add_para(doc, 'The procedure for each received frame:')
    for i, step in enumerate([
        'Monitor rx_sync in the idle state. A falling edge (logic 0) indicates '
            'the start of the start bit. Reset the tick counter.',
        'Count 7 s_ticks — this positions sampling at the midpoint of the start bit '
            '(tick 7 of 16). Confirm the line is still low to reject noise glitches.',
        'Count 16 s_ticks — now at the midpoint of data bit D0 (tick 15). '
            'Sample rx_sync and shift into the data register.',
        'Repeat step 3 for the remaining 7 data bits.',
        'If parity is enabled, repeat step 3 once to sample the parity bit.',
        'Count through the stop bit period; verify rx_sync is high. '
            'Assert rx_done_tick.',
    ], 1):
        add_bullet(doc, f'{i}.  {step}')

    add_figure(doc, FIGPATHS['oversampling'],
               'Figure 2.  16× Oversampling — Start Bit and First Data Bit',
               width=Inches(6.0))

    add_heading(doc, 'Input Synchronization (2-FF Synchronizer)', 2)
    add_para(doc,
        'The rx line arrives from an external device and is therefore asynchronous '
        'to the FPGA\'s clock domain. Sampling it directly risks metastability — '
        'a flip-flop entering an undefined voltage state that can propagate corrupted '
        'values through the logic.')
    add_para(doc,
        'The solution is a two flip-flop (2-FF) synchronizer placed at the RX input '
        'before the FSM. The first FF may enter a metastable state, but it has a '
        'full clock period to resolve before the second FF (rx_sync) samples it. '
        'Both FFs reset to logic 1 (the UART idle level) to prevent a false '
        'start-bit detection on power-up.')

    add_figure(doc, FIGPATHS['2ff'],
               'Figure 3.  2-FF Metastability Synchronizer on the RX Input',
               width=Inches(5.5))

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    #  4. TOP-LEVEL DESIGN
    # ════════════════════════════════════════════════════════════
    add_heading(doc, '4.  Top-Level Design Diagram', 1)
    add_para(doc,
        'The system is composed of three primary blocks instantiated inside '
        'UART_top: the Baud Rate Generator, the Transmitter, and the Receiver. '
        'The Baud Rate Generator\'s done output is distributed to both TX and RX '
        'as the shared s_tick timing reference. The parity_select input is also '
        'shared, ensuring both sides always use the same parity configuration.')

    add_figure(doc, FIGPATHS['toplevel'],
               'Figure 4.  UART_top — Top-Level Block Diagram',
               width=Inches(6.0))

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    #  5. COMPONENT BREAKDOWN
    # ════════════════════════════════════════════════════════════
    add_heading(doc, '5.  Component Breakdown', 1)

    # 5.1 BRG
    add_heading(doc, '5.1  Baud Rate Generator  —  Baud_Rate_Generator.v', 2)
    add_para(doc,
        'A parameterized modulo-N counter. It divides the system clock to produce '
        'a single-cycle pulse (done) at 16× the target baud rate. The counter '
        'increments on every rising clock edge while en is high. When the count '
        'reaches FINAL_VALUE, done is asserted for exactly one clock cycle and '
        'the counter resets to zero.')

    add_heading(doc, 'Parameters', 3)
    add_table(doc,
        ['Parameter', 'Default', 'Description'],
        [('BITS', '4',
          'Width of the internal counter and the FINAL_VALUE input')],
        col_widths=[Cm(3), Cm(2.5), Cm(9.5)])

    add_heading(doc, 'Ports', 3)
    add_table(doc,
        ['Port', 'Direction', 'Description'],
        [('clk',         'Input',  'System clock'),
         ('reset_n',     'Input',  'Active-low asynchronous reset'),
         ('en',          'Input',  'Counter enable — holds count when low'),
         ('FINAL_VALUE', 'Input',  '[BITS-1:0] — counter rolls over at this value'),
         ('done',        'Output', 'High for one clock cycle when count == FINAL_VALUE')],
        col_widths=[Cm(3.2), Cm(2.5), Cm(9.3)])

    # 5.2 TX
    add_heading(doc, '5.2  UART Transmitter  —  UART_tx.v', 2)
    add_para(doc,
        'Converts an 8-bit parallel word into a serial UART frame. Controlled by '
        'a 5-state Algorithmic State Machine with Datapath (ASMD). The parity bit '
        'is pre-computed in the idle state when tx_start is asserted — before '
        'data shifting begins — and stored in p_reg. This is necessary because '
        'the shift register b_reg is destroyed during the data state.')

    add_heading(doc, 'Parameters', 3)
    add_table(doc,
        ['Parameter', 'Default', 'Description'],
        [('DBITS',   '8',  'Number of data bits per frame'),
         ('SB_TICK', '16', 'Stop-bit duration in s_ticks (16=1 stop bit, 32=2 stop bits)')],
        col_widths=[Cm(3), Cm(2.5), Cm(9.5)])

    add_heading(doc, 'Ports', 3)
    add_table(doc,
        ['Port', 'Direction', 'Description'],
        [('clk',           'Input',  'System clock'),
         ('reset_n',       'Input',  'Active-low asynchronous reset'),
         ('s_tick',        'Input',  '16× baud rate sampling tick from BRG'),
         ('tx_start',      'Input',  'Assert high for one cycle to initiate transmission'),
         ('parity_select', 'Input',  "[1:0] — 2'b00=none, 2'b01=even, 2'b10=odd"),
         ('tx_din',        'Input',  '[DBITS-1:0] — parallel data word to transmit'),
         ('tx_done_tick',  'Output', 'Pulses high for one cycle when frame is complete'),
         ('tx',            'Output', 'Serial data output line')],
        col_widths=[Cm(3.5), Cm(2.5), Cm(9.0)])

    add_heading(doc, 'State Machine', 3)
    add_table(doc,
        ['State', 'TX Line', 'Action'],
        [('idle',   '1',         'Wait for tx_start. On assertion: load tx_din, pre-compute parity into p_reg.'),
         ('start',  '0',         'Drive start bit low for 16 s_ticks, then advance to data.'),
         ('data',   'b_reg[0]',  'Shift out data bits LSB-first, one bit per 16 s_ticks. After last bit advance to parity or stop.'),
         ('parity', 'p_reg',     'Transmit pre-computed parity bit for 16 s_ticks, then advance to stop.'),
         ('stop',   '1',         'Drive stop bit(s) for SB_TICK s_ticks. Assert tx_done_tick, return to idle.')],
        col_widths=[Cm(2.5), Cm(2.5), Cm(10.0)])

    # 5.3 RX
    add_heading(doc, '5.3  UART Receiver  —  UART_rx.v', 2)
    add_para(doc,
        'Recovers an 8-bit parallel word from an incoming serial UART frame '
        'using 16× oversampling. Includes a 2-FF input synchronizer for the raw '
        'rx input, parity checking, and frame error detection. All FSM logic '
        'operates on the synchronised signal rx_sync.')

    add_heading(doc, 'Parameters', 3)
    add_table(doc,
        ['Parameter', 'Default', 'Description'],
        [('DBITS',   '8',  'Number of data bits per frame'),
         ('SB_TICK', '16', 'Stop-bit duration in s_ticks')],
        col_widths=[Cm(3), Cm(2.5), Cm(9.5)])

    add_heading(doc, 'Ports', 3)
    add_table(doc,
        ['Port', 'Direction', 'Description'],
        [('clk',           'Input',  'System clock'),
         ('reset_n',       'Input',  'Active-low asynchronous reset'),
         ('rx',            'Input',  'Raw serial input — passes through 2-FF synchronizer'),
         ('s_tick',        'Input',  '16× baud rate sampling tick from BRG'),
         ('parity_select', 'Input',  "[1:0] — 2'b00=none, 2'b01=even, 2'b10=odd"),
         ('rx_dout',       'Output', '[DBITS-1:0] — recovered parallel data word'),
         ('rx_done_tick',  'Output', 'Pulses high for one cycle when a complete frame is received'),
         ('parity_error',  'Output', 'Pulses high (same cycle as rx_done_tick) on parity mismatch'),
         ('frame_error',   'Output', 'Pulses high (same cycle as rx_done_tick) when stop bit is logic 0')],
        col_widths=[Cm(3.5), Cm(2.5), Cm(9.0)])

    add_heading(doc, 'State Machine', 3)
    add_table(doc,
        ['State', 'Advance Condition', 'Action'],
        [('idle',   'rx_sync == 0',          'Start bit detected. Reset tick counter, advance to start.'),
         ('start',  's_tick && s_reg == 7',   'Confirm start bit at midpoint. Reset counters, advance to data.'),
         ('data',   's_tick && s_reg == 15',  'Sample rx_sync, right-shift into b_reg. Repeat for all DBITS. Advance to parity or stop.'),
         ('parity', 's_tick && s_reg == 15',  'Latch rx_sync into p_reg. Advance to stop.'),
         ('stop',   's_tick && s_reg == SB_TICK-1', 'Check stop bit and parity. Assert rx_done_tick, return to idle.')],
        col_widths=[Cm(2.5), Cm(4.2), Cm(8.3)])

    add_heading(doc, 'Error Detection Logic', 3)
    add_para(doc,
        'Both error signals are combinational and valid on the same clock cycle '
        'as rx_done_tick:')
    for txt in [
        'frame_error  = ~rx_sync                          (stop bit must be logic 1)',
        'parity_error = (^b_reg) != p_reg                 [even parity]',
        'parity_error = (~^b_reg) != p_reg                [odd  parity]',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(txt)
        r.font.name = 'Courier New'; r.font.size = Pt(9.5)

    # 5.4 Top
    add_heading(doc, '5.4  Top-Level Wrapper  —  UART_top.v', 2)
    add_para(doc,
        'Integrates the three sub-modules. The BRG output s_tick is shared '
        'between TX and RX. The parity_select input is also shared, ensuring '
        'both sides always use the same parity configuration.')

    add_heading(doc, 'Parameters', 3)
    add_table(doc,
        ['Parameter', 'Default', 'Description'],
        [('DBITS',     '8',  'Data bits per frame'),
         ('SB_TICK',   '16', 'Stop-bit duration in s_ticks'),
         ('BRG_BITS',  '4',  'Counter width of the Baud Rate Generator'),
         ('BRG_FINAL', '7',  'Counter limit — set per baud rate formula in Section 2')],
        col_widths=[Cm(3.2), Cm(2.5), Cm(9.3)])

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    #  6. ASMD CHARTS
    # ════════════════════════════════════════════════════════════
    add_heading(doc, '6.  ASMD Charts', 1)

    add_heading(doc, 'ASMD Chart of the UART Transmitter', 2)
    add_para(doc,
        'The transmitter ASMD chart shows the five-state FSM and its datapath '
        'operations. State boxes (dashed blue borders) show Moore outputs. '
        'Decision diamonds show Mealy conditions. Operation boxes (green) show '
        'datapath register updates.')
    add_figure(doc, FIGPATHS['tx_asmd'],
               'Figure 5.  UART Transmitter ASMD Chart',
               width=Inches(4.8))

    page_break(doc)

    add_heading(doc, 'ASMD Chart of the UART Receiver', 2)
    add_para(doc,
        'The receiver ASMD chart mirrors the transmitter structure with the addition '
        'of the parity and frame error checks in the stop state. The rx_done_tick, '
        'parity_error, and frame_error signals are all asserted combinationally '
        'within the stop state when s_reg == SB_TICK-1.')
    add_figure(doc, FIGPATHS['rx_asmd'],
               'Figure 6.  UART Receiver ASMD Chart',
               width=Inches(4.8))

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    #  7. TESTBENCHES & SIMULATION
    # ════════════════════════════════════════════════════════════
    add_heading(doc, '7.  Testbenches & Simulation', 1)

    add_heading(doc, '7.1  Transmitter Testbench  —  UART_tx_tb.v', 2)
    add_para(doc,
        'Tests the transmitter module in isolation. The Baud Rate Generator runs '
        'with FINAL_VALUE = 7 (s_tick every 8 clock cycles) so that the individual '
        'bits on the tx line are clearly visible in the waveform. The testbench '
        'sends 7 bytes across three parity configurations using a send_byte task '
        'that asserts tx_start for one clock cycle and waits for tx_done_tick.')

    add_table(doc,
        ['#', 'Data', 'Format', 'Description'],
        [('1', '0xE5', '8N1', '8 data bits, no parity, 1 stop bit'),
         ('2', '0x3C', '8N1', '8 data bits, no parity, 1 stop bit'),
         ('3', '0xF0', '8N1', '8 data bits, no parity, 1 stop bit'),
         ('4', '0xA5', '8E1', '8 data bits, even parity, 1 stop bit'),
         ('5', '0x55', '8E1', '8 data bits, even parity — alternating bits'),
         ('6', '0xAA', '8O1', '8 data bits, odd parity  — alternating bits'),
         ('7', '0xFF', '8O1', '8 data bits, odd parity  — all ones')],
        col_widths=[Cm(1.0), Cm(2.2), Cm(2.0), Cm(9.8)])

    add_para(doc,
        'The simulation waveform shows one complete byte transmission. The tx signal '
        'is driven low for the start bit, transitions between logic levels '
        'corresponding to each data bit (LSB first), then shows the parity bit '
        'and stop bit. Each level is held for exactly 16 s_ticks (one bit period).')

    add_figure(doc,
               os.path.join(SCREENDIR, 'tx_tb_one_byte.png'),
               'Figure 7.  UART Transmitter Simulation Waveform — single byte transmission showing '
               'start bit, data bits, parity bit, and stop bit on the tx line',
               width=Inches(6.0))

    page_break(doc)

    add_heading(doc, '7.2  Full UART Loopback Testbench  —  UART_loopback_tb.v', 2)
    add_para(doc,
        'Connects the transmitter\'s tx output directly to the receiver\'s rx input, '
        'testing the complete data path end-to-end. The BRG runs with FINAL_VALUE = 0 '
        '(s_tick fires every clock cycle) to maximise simulation speed. A '
        'send_and_verify task transmits each byte and automatically checks rx_dout, '
        'parity_error, and frame_error, printing PASS or FAIL to the console.')

    add_table(doc,
        ['#', 'Data', 'Format', 'Expected Result'],
        [('1', '0xA5', '8N1',  'rx_dout = 0xA5, no errors'),
         ('2', '0x55', '8E1',  'rx_dout = 0x55, no errors'),
         ('3', '0xAA', '8O1',  'rx_dout = 0xAA, no errors'),
         ('4', '0xF0', '8N1',  'rx_dout = 0xF0, no errors'),
         ('5', '0x3C', '8E1',  'rx_dout = 0x3C, no errors'),
         ('6', '0xE5', '8O1',  'rx_dout = 0xE5, no errors')],
        col_widths=[Cm(1.0), Cm(2.2), Cm(2.0), Cm(9.8)])

    add_para(doc,
        'The simulation results confirm that all 6 test cases pass. The data '
        'recovered by the receiver (rx_dout) exactly matches the data loaded into '
        'the transmitter (tx_din) in every case. The parity_error and frame_error '
        'signals remain deasserted throughout, confirming correct parity generation '
        'and frame delimitation across all three parity configurations.')

    add_figure(doc,
               os.path.join(SCREENDIR, 'loopbakc_zoomed_out.png'),
               'Figure 8.  UART Loopback Simulation — all 6 test cases. tx_din and rx_dout '
               'match in every frame; parity_error and frame_error remain at 0.',
               width=Inches(6.0))

    add_figure(doc,
               os.path.join(SCREENDIR, 'loopback_one_byte.png'),
               'Figure 9.  UART Loopback Simulation — single byte zoomed in. The RX FSM '
               'state transitions and the b_reg shift register building up bit by bit are visible.',
               width=Inches(6.0))

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    #  8. REFERENCES
    # ════════════════════════════════════════════════════════════
    add_heading(doc, '8.  References', 1)
    for ref in [
        'Chu, P. P. (2008). FPGA Prototyping by Verilog Examples: Xilinx Spartan-3 '
            'Version. Wiley-Interscience.',
        'Anas Salah Eddin. ECE 3300 — Digital Circuit Design Using Verilog. '
            'YouTube playlist.',
        'Xilinx / AMD. Vivado Design Suite User Guide: Logic Simulation (UG900). '
            'AMD, 2023.',
        'IEEE Standard for Verilog Hardware Description Language. '
            'IEEE Std 1364-2005. IEEE, 2006.',
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent  = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        r = p.add_run(ref)
        r.font.size = Pt(10.5)

    # ── Save ─────────────────────────────────────────────────────
    out = os.path.join(BASE, 'UART_Documentation.docx')
    doc.save(out)
    print(f'Saved: {out}')


# ══════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print('Generating figures...')
    fig_frame()
    fig_oversampling()
    fig_2ff()
    fig_toplevel()
    fig_tx_asmd()
    fig_rx_asmd()
    print(f'  {len(FIGPATHS)} figures created.')
    print('Building Word document...')
    build_doc()
    print('Done.')
